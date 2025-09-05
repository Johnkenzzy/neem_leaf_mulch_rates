#!/usr/bin/python3
import argparse
import pandas as pd
import statsmodels.api as sm
from statsmodels.formula.api import ols
import scikit_posthocs as sp
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os
import warnings

warnings.filterwarnings("ignore", message="Precision loss occurred in moment calculation")

# ---------- Helper Function to Assign Letters ----------
def assign_letters(means, pvals, alpha=0.05):
    letters = {t: "" for t in means.index}
    sorted_means = means.sort_values(ascending=False)
    current_letter = "a"

    for treatment in sorted_means.index:
        assigned = False
        for existing_letter in set(letters.values()):
            if not existing_letter:
                continue
            same_group = True
            for t2, letter in letters.items():
                if letter == existing_letter:
                    p = pvals.loc[treatment, t2] if (treatment in pvals.index and t2 in pvals.columns) else 1
                    if p < alpha:
                        same_group = False
                        break
            if same_group:
                letters[treatment] = existing_letter
                assigned = True
                break
        if not assigned:
            letters[treatment] = current_letter
            current_letter = chr(ord(current_letter) + 1)
    return letters

# ---------- Main Function ----------
def analyze_data(file_path, sheet_name):
    # Load Excel sheet
    df = pd.read_excel(file_path, sheet_name=sheet_name)

    # --- Dynamically detect treatment & numeric columns ---
    treatment_col = None
    for candidate in ["mulch_rate", "treatment", "trt"]:
        if candidate in df.columns:
            treatment_col = candidate
            break
    if not treatment_col:
        raise ValueError("No treatment column found (expected: mulch_rate, treatment, trt)")

    # Get all numeric columns except treatment
    response_vars = df.select_dtypes(include="number").columns.tolist()
    if treatment_col in response_vars:
        response_vars.remove(treatment_col)

    if not response_vars:
        raise ValueError("No numeric columns found for ANOVA.")

    all_results = []
    doc = Document()
    doc.add_heading("Crop Experiment ANOVA & LSD Report", 0)

    for var in response_vars:
         # Drop missing rows for treatment and this variable
        df_var = df[[treatment_col, var]].dropna()
        if df_var.empty:
            print(f"\n⚠️ Skipping {var} (no valid data)")
            continue
        print(f"\n=== ANOVA for {var} ===")
        model = ols(f'Q("{var}") ~ C({treatment_col})', data=df).fit()
        anova = sm.stats.anova_lm(model, typ=2).fillna("-")
        print(anova)

        # --- Add ANOVA table to Word ---
        doc.add_heading(var, level=1)
        doc.add_paragraph("ANOVA Table:")
        table = doc.add_table(rows=1, cols=len(anova.columns) + 1)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Source"
        for i, col in enumerate(anova.columns):
            hdr_cells[i+1].text = col
        for idx, row in anova.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            for i, val in enumerate(row):
                row_cells[i+1].text = f"{val:.4f}" if isinstance(val, (float, int)) else str(val)

        means = df.groupby(treatment_col)[var].mean()

        if isinstance(anova['PR(>F)'].iloc[0], float) and anova['PR(>F)'].iloc[0] < 0.05:
            print("Significant effect found → running LSD test")
            lsd = sp.posthoc_ttest(df_var, val_col=var, group_col=treatment_col, p_adjust=None)
            letters = assign_letters(means, lsd)

            # LSD results → Word table
            doc.add_paragraph("LSD Grouping:")
            lsd_table = doc.add_table(rows=1, cols=3)
            hdr_cells = lsd_table.rows[0].cells
            hdr_cells[0].text = treatment_col
            hdr_cells[1].text = "Mean"
            hdr_cells[2].text = "Group"
            for t in means.index:
                all_results.append([var, t, means[t], letters[t]])
                row_cells = lsd_table.add_row().cells
                row_cells[0].text = str(t)
                row_cells[1].text = f"{means[t]:.2f}"
                row_cells[2].text = letters[t]

            # Plot + annotate
            plt.figure(figsize=(7, 4))
            means.plot(kind='bar', yerr=df.groupby(treatment_col)[var].std(), capsize=4)
            for i, (t, mean) in enumerate(means.items()):
                plt.text(i, mean + 0.05 * mean, letters[t], ha='center', va='bottom',
                         fontsize=12, fontweight='bold')
            plt.title(f"{var} by {treatment_col}")
            plt.ylabel(var)
            plt.xlabel(f"{treatment_col}(kg/ha)")
            plt.tight_layout()

            plot_path = os.path.join(os.path.dirname(file_path), f"{var}_plot.png")
            plt.savefig(plot_path)
            plt.close()
            doc.add_picture(plot_path, width=Inches(5))
        else:
            print("No significant difference between treatments.")
            for t in means.index:
                all_results.append([var, t, means[t], "-"])

    # --- Save results ---
    output_xlsx = os.path.join(os.path.dirname(file_path), "analysis_results.xlsx")
    results_df = pd.DataFrame(all_results, columns=["variable", treatment_col, "mean", "group_letter"])
    results_df.to_excel(output_xlsx, index=False)

    output_docx = os.path.join(os.path.dirname(file_path), "analysis_report.docx")
    doc.save(output_docx)

    print(f"\n✅ Results saved to:\n   Excel: {output_xlsx}\n   Word: {output_docx}")

# ---------- CLI ----------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Analyze crop experiment data (ANOVA + LSD)")
    parser.add_argument("file", help="Path to the Excel file")
    parser.add_argument("sheet", help="Sheet name to analyze")
    args = parser.parse_args()

    analyze_data(args.file, args.sheet)
